def windows_print_file(fp,printer):
    '''window系统文件打印'''
    try:
        import win32api
        import win32print
        # 设置默认打印机
        win32print.SetDefaultPrinter(printer)
        win32api.ShellExecute( 0,"print",fp,'/d:"%s"' % win32print.GetDefaultPrinter(),".", 0)
        return True
    except Exception as e:
        import traceback
        traceback.print_exc()
        return False

def saveAsDocx(file):
    '''保存文件未docx'''
    try:
        import win32com.client
        word = win32com.client.gencache.EnsureDispatch('Word.Application')
        word.Visible = False
        word.DisplayAlerts = False
        doc = word.Documents.Open(file, False)  # 打开文档，不提示转换确认框
        new_file = file.split('.')[0] + '.docx'
        doc.SaveAs(new_file, 12)
        doc.Close()
        return new_file
    except Exception as e:
        import traceback
        traceback.print_exc()
        return None

def combineAsDocx(master, sub):
    '''合并docx'''
    try:
        import os
        import docx
        import docxcompose.composer
        # 待合并文件必须存在
        if not os.path.exists(sub):
            return False

        # 主文件必须是docx格式,可以不存在
        if not master.endswith('.docx'):
            return False

        sub_docx = sub
        if not sub.endswith('.docx'):
            sub_docx = saveAsDocx(sub)

        if os.path.exists(master):
            doc_master = docx.Document(master)

            cp = docxcompose.composer.Composer(doc_master)
            cp.append(docx.Document(sub_docx))
        else:
            # master不存在，则sub直接给master
            doc_master = docx.Document(sub_docx)
        doc_master.save(master)
        return True
    except Exception as e:
        import traceback
        traceback.print_exc()
        return False